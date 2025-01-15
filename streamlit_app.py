import nltk
import requests
import os
import tempfile
import streamlit as st
import time
import openai
import langchain 
from langchain.llms import OpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import FAISS 
from langchain.schema import Document 
from langchain.chat_models import ChatOpenAI
from langchain.chains import ConversationalRetrievalChain 
import io
from PyPDF2 import PdfReader
from docx import Document as DocxDocument 


# Get all environment variables
from dotenv import load_dotenv
load_dotenv()
openai.api_key = st.secrets["OPENAI_API_KEY"]


# Web interface
st.title("DataScout AI")
st.sidebar.subheader("Websites")

urls = []
for i in range(3):
    url = st.sidebar.text_input(f"URL {i+1}")
    urls.append(url)

# File upload 
st.sidebar.subheader('Upload documents')
uploaded_files = st.sidebar.file_uploader("Select files from your computer", type=['txt','pdf','docx'], accept_multiple_files=True)

process_url_clicked = st.sidebar.button("Get content")

# progress bar
main_placefolder = st.empty()


# Initialize llm 
llm = ChatOpenAI(temperature=0.6, max_tokens=500, model="gpt-3.5-turbo")

if process_url_clicked:
    # load data from URLs
    data = []
    if any(urls):
        for url in urls:
            if url.strip():  # Ensure URL is not empty
                try:
                    main_placefolder.text(f"Loading content from: {url} . . . ⏳")
                    
                    # Fetch content using requests
                    response = requests.get(url.strip(), timeout=20)  # 20-second timeout
                    if response.status_code == 200:
                        main_placefolder.text(f"Successfully fetched content from: {url}")
                        content = response.text  # Extract HTML or plain text
                        
                        # Append as a dictionary 
                        data.append({"url": url, "content": content})
                    else:
                        st.error(f"Failed to fetch URL {url}. Status: {response.status_code}")
                except Exception as e:
                    st.error(f"Error fetching URL {url}: {e}")

    
    # load data from uploaded documents
    if any(uploaded_files):
        for uploaded_file in uploaded_files:
            if uploaded_file.name.strip(): # ensure file name is not empty
                try:
                    main_placefolder.text(f"Processing file: {uploaded_file.name} . . .⏳ ")

                    # validate file extensions
                    valid_extensions = (".pdf",".docx",".txt")
                    if not uploaded_file.name.lower().endswith(valid_extensions):
                        st.warning(f"Skipping unsupported file type: {uploaded_file.name}")
                        continue
                    
                    # Read and process file content based on file type
                    if uploaded_file.name.lower().endswith(".pdf"):
                        pdf_reader = PdfReader(io.BytesIO(uploaded_file.read()))
                        file_content = "\n".join(page.extract_text() for page in pdf_reader.pages)
                    
                    elif uploaded_file.name.lower().endswith(".docx"):
                        docx = DocxDocument(io.BytesIO(uploaded_file.read()))
                        file_content = "\n".join(paragraph.text for paragraph in docx.paragraphs)
                    
                    elif uploaded_file.name.lower().endswith(".txt"):
                        file_content = uploaded_file.read().decode("utf-8", errors="ignore")
                    
                    else:
                        st.error(f"Unsupported file format: {uploaded_file.name}")
                        continue

                    # append dictionary for file name and content
                    data.append({"file_name":uploaded_file.name, "content":file_content})

                    main_placefolder.text(f"Successfully loaded content from {uploaded_file.name}")
                
                except Exception as e:
                    st.error(f"Error processing file {uploaded_file.name}. \nError: {e}")

    

    # check if data was loaded properly
    if not data:
        st.error("No data was loaded from the provided URLs or files. Please check the inputs")
        st.stop()
    else:
        st.write(f"Loaded {len(data)} document(s)")
    
    # Convert each dictionary into a Document object
    data = [
        Document(
            page_content=doc['content'], 
            metadata={"url": doc['url']} if "url" in doc else {"file_name":doc['file_name']}
        ) 
        for doc in data
    ]

    # split data
    text_splitter = RecursiveCharacterTextSplitter(
        separators=['\n\n', '\n', '.', ','],
        chunk_size=1000
    )
    main_placefolder.text("Splitting your data. . . . ✂")
    docs = text_splitter.split_documents(data)

    # validate text splitting
    if not docs:
        st.error("No documents were created after splitting. Ensure the data contains valid text")
        st.stop()
    else:
        st.write(f"Split into {len(docs)} chunks")

    # embeddings
    # store embeddings in FAISS index 
    # and FAISS index in Streamlit session state memory
    embeddings = OpenAIEmbeddings()
    vectorstore_openai = FAISS.from_documents(docs, embeddings)
    st.session_state['vectorstore'] = vectorstore_openai
    main_placefolder.text("Creating embedding vector. . . . 📚")
    time.sleep(1)

    # Reset trigger
    process_url_clicked = False

# Question box and submit button
with st.form("query_form"):
    query = st.text_input("What do you want to know?", "")
    enter_button = st.form_submit_button("Enter")

if enter_button:
    if query:
        # Load FAISS index from session state
        if 'vectorstore' in st.session_state:
            vectorstore = st.session_state['vectorstore']

            # Initialize chat_history as an empty list or use previous conversation history
            chat_history = []

            # Create the Conversational Retrieval Chain
            conversation_chain = ConversationalRetrievalChain.from_llm(llm=llm, retriever=vectorstore.as_retriever())

            # Query the chain
            result = conversation_chain({"question": query, "chat_history": chat_history})

            # Display the answer
            st.header("Answer")
            st.write(result["answer"])

            
        else: 
            st.error("FAISS index not found. Please process data first")

